"""
知识库加载器 - 从docx/pdf文件加载并处理知识库
"""
import os
import re
import logging
from typing import List, Dict
import PyPDF2
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.schema import Document as LangchainDocument

logger = logging.getLogger(__name__)


class KnowledgeLoader:
    """知识库加载和管理类"""

    def __init__(self, persist_directory="./chroma_db"):
        """
        初始化知识库加载器

        Args:
            persist_directory: Chroma数据库持久化目录
        """
        self.persist_directory = persist_directory

        # 使用text2vec-base-chinese模型进行embedding
        logger.info("初始化text2vec-base-chinese模型...")
        self.embeddings = HuggingFaceEmbeddings(
            model_name="shibing624/text2vec-base-chinese",
            model_kwargs={'device': 'cpu'},
            encode_kwargs={'normalize_embeddings': True}
        )

        self.vectorstore = None
        self.documents = []

    def load_docx(self, file_path: str) -> str:
        """
        加载 docx 文件内容，优先使用 python-docx，
        如果失败则回退到 docx2txt。
        返回一个包含所有段落的字符串。
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        try:
            from docx import Document
            doc = Document(file_path)
            full_text = [p.text for p in doc.paragraphs]  # 不去掉空行
            return "\n".join(full_text)
        except Exception as e:
            try:
                import docx2txt
                return docx2txt.process(file_path)
            except Exception as e2:
                raise RuntimeError(
                    "读取 docx 文件失败：\n"
                    "1) 确认已安装 python-docx： python -m pip install python-docx\n"
                    "2) 或安装 docx2txt 作为备选： python -m pip install docx2txt\n"
                    f"原始错误: {e}\n回退错误: {e2}"
                )

    def load_pdf(self, file_path: str) -> str:
        """
        加载PDF文件
        """
        try:
            logger.info(f"加载PDF文件: {file_path}")
            full_text = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text and text.strip():
                        full_text.append(text.strip())
            content = '\n'.join(full_text)
            logger.info(f"成功加载PDF文件，共 {len(pdf_reader.pages)} 页")
            return content
        except Exception as e:
            logger.error(f"加载PDF文件失败: {e}")
            raise

    def split_by_empty_lines(self, text: str) -> List[str]:
        """
        按空行分块文本
        - 空行 = 至少一个空白行
        - 每个块保留原段落（不逐行拼接）
        """
        # 使用正则：一个或多个换行+空白 => 分隔符
        raw_parts = re.split(r"\n\s*\n", text)
        # 去掉首尾空格和空块
        chunks = [part.strip() for part in raw_parts if part.strip()]
        logger.info(f"文本分块完成，共 {len(chunks)} 个块")
        return chunks

    def create_documents(self, chunks: List[str], source: str = "knowledge_base") -> List[LangchainDocument]:
        """将文本块转换为Langchain Document对象"""
        documents = []
        for i, chunk in enumerate(chunks):
            metadata = {
                "source": source,
                "chunk_id": i,
                "chunk_length": len(chunk)
            }
            doc = LangchainDocument(
                page_content=chunk,
                metadata=metadata
            )
            documents.append(doc)
        logger.info(f"创建了 {len(documents)} 个Document对象")
        return documents

    def build_vectorstore(self, documents: List[LangchainDocument]) -> Chroma:
        """构建Chroma向量数据库"""
        try:
            logger.info("开始构建Chroma向量数据库...")
            self.vectorstore = Chroma.from_documents(
                documents=documents,
                embedding=self.embeddings,
                persist_directory=self.persist_directory
            )
            self.vectorstore.persist()
            logger.info(f"Chroma向量数据库构建完成，共 {len(documents)} 个文档")
            return self.vectorstore
        except Exception as e:
            logger.error(f"构建向量数据库失败: {e}")
            raise

    def load_vectorstore(self) -> Chroma:
        """加载已存在的Chroma向量数据库"""
        try:
            if not os.path.exists(self.persist_directory):
                raise FileNotFoundError(f"向量数据库目录不存在: {self.persist_directory}")
            logger.info(f"加载现有Chroma向量数据库: {self.persist_directory}")
            self.vectorstore = Chroma(
                persist_directory=self.persist_directory,
                embedding_function=self.embeddings
            )
            logger.info("成功加载向量数据库")
            return self.vectorstore
        except Exception as e:
            logger.error(f"加载向量数据库失败: {e}")
            raise

    def process_knowledge_file(self, file_path: str, force_rebuild: bool = False):
        """
        处理知识库文件的完整流程
        """
        if os.path.exists(self.persist_directory) and not force_rebuild:
            logger.info("检测到现有数据库，直接加载...")
            self.load_vectorstore()
            return

        file_ext = os.path.splitext(file_path)[1].lower()
        if file_ext == '.docx':
            content = self.load_docx(file_path)
        elif file_ext == '.pdf':
            content = self.load_pdf(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {file_ext}")

        # 2. 按空行分块
        chunks = self.split_by_empty_lines(content)

        # 3. 创建Document对象
        self.documents = self.create_documents(
            chunks,
            source=os.path.basename(file_path)
        )

        # 4. 构建向量数据库
        self.build_vectorstore(self.documents)
        logger.info("知识库处理完成！")

    def search(self, query: str, k: int = 5) -> List[Dict]:
        """搜索相关文档"""
        if not self.vectorstore:
            raise ValueError("向量数据库未初始化，请先加载或构建数据库")

        try:
            results = self.vectorstore.similarity_search_with_score(query, k=k)
            formatted_results = []
            for doc, score in results:
                formatted_results.append({
                    "content": doc.page_content,
                    "metadata": doc.metadata,
                    "similarity_score": float(score)
                })
            logger.info(f"搜索完成，找到 {len(formatted_results)} 个相关结果")
            return formatted_results
        except Exception as e:
            logger.error(f"搜索失败: {e}")
            raise

    def get_stats(self) -> Dict:
        """获取知识库统计信息"""
        if not self.vectorstore:
            return {"status": "未初始化"}
        try:
            collection = self.vectorstore._collection
            count = collection.count()
            return {
                "status": "已加载",
                "total_documents": count,
                "persist_directory": self.persist_directory,
                "embedding_model": "text2vec-base-chinese"
            }
        except Exception as e:
            logger.error(f"获取统计信息失败: {e}")
            return {"status": "错误", "error": str(e)}
